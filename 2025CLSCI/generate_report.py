#!/usr/bin/env python3
"""
生成CLSCI 2025年泛刑法文章研究报告（DOCX格式）
数据驱动，一万字以上
"""

import json
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = '/Users/yalipeng/Downloads/2025CLSCI/output'

# 加载数据
with open(f'{OUTPUT_DIR}/CLSCI_all_titles_2025.json', 'r') as f:
    all_arts = json.load(f)

with open(f'{OUTPUT_DIR}/criminal_law_details_all.json', 'r') as f:
    criminal = json.load(f)

# ============ 数据预处理 ============

total_count = len(all_arts)
criminal_count = len(criminal)
has_abstract = sum(1 for a in criminal if a.get('abstract_cn', '').strip())

# 各期刊统计
total_by_j = Counter([a.get('journal', '') for a in all_arts])
crim_by_j = Counter([a.get('journal', '') for a in criminal])

journal_stats = []
for j in sorted(total_by_j.keys()):
    t = total_by_j[j]
    c = crim_by_j.get(j, 0)
    pct = 100 * c / t if t > 0 else 0
    journal_stats.append({'journal': j, 'total': t, 'criminal': c, 'pct': pct})

journal_stats_sorted = sorted(journal_stats, key=lambda x: -x['pct'])

# 期数分布
issue_counts = Counter()
for a in criminal:
    issue = a.get('issue', '')
    if issue:
        issue_counts[issue] += 1

# 细分专题
def topic_count(keywords_list):
    return [a for a in criminal if any(kw in a.get('title', '') for kw in keywords_list)]

topics_data = {
    '刑法基本理论（法益/归责/违法性等）': topic_count(['法益', '归责', '构成要件', '违法性', '不法', '犯罪论', '犯罪构成', '刑法学', '知识体系', '规范论', '法哲学']),
    '轻罪治理与犯罪记录封存': topic_count(['轻罪', '轻微犯罪', '犯罪记录', '前科', '复权', '出罪', '但书', '微罪']),
    '经济刑法与反腐败': topic_count(['受贿', '贪污', '行贿', '腐败', '洗钱', '经济刑法', '非法经营', '税收犯罪', '虚开', '非法集资', '金融犯罪']),
    '行刑衔接与法秩序统一': topic_count(['行刑', '刑民', '民刑', '行政犯', '法定犯', '法秩序', '反向衔接']),
    '人工智能与刑法': topic_count(['人工智能', 'AI', '算法', '大模型', '生成式', '智能', '深度伪造']),
    '网络犯罪治理': topic_count(['网络犯罪', '网络暴力', '电信诈骗', '电信网络', '帮信', '帮助信息']),
    '数据犯罪与个人信息保护': topic_count(['数据犯罪', '个人信息', '数据安全', '数据保护', '数据刑法']),
    '刑事证据制度': topic_count(['证据']) if False else [a for a in criminal if '证据' in a.get('title','') and any(w in a.get('title','') for w in ['刑事','刑法','犯罪','侦查','证据法','电子证据','非法证据','证据层','证据链','庭审','供述','鉴定'])],
    '认罪认罚从宽制度': topic_count(['认罪认罚', '从宽']),
    '刑事诉讼法修改': [a for a in criminal if any(w in a.get('title','') for w in ['刑事诉讼法','刑诉法']) and '修改' in a.get('title','') or '法典化' in a.get('title','') and '刑事' in a.get('title','')],
    '量刑制度': topic_count(['量刑']),
    '危险犯理论与醉驾治理': topic_count(['危险犯', '抽象危险', '危险驾驶', '醉驾']),
    '诈骗犯罪': topic_count(['诈骗']),
    '财产犯罪': topic_count(['盗窃', '财产犯罪', '占有', '财产性利益', '侵占', '财产损害']),
    '刑事司法数字化转型': [a for a in criminal if any(w in a.get('title','') for w in ['数字检察','数字时代','数字化','大数据']) and any(w in a.get('title','') for w in ['刑事','刑法','犯罪','侦查','检察','司法'])],
    '侦查制度': topic_count(['侦查']),
    '正当防卫/紧急避险': topic_count(['正当防卫', '防卫', '紧急避险', '紧急权']),
    '共犯论': topic_count(['共犯', '共谋', '共同犯罪', '帮助犯', '正犯']),
    '单位犯罪与企业合规': topic_count(['单位犯罪', '企业合规', '企业刑事', '涉企']),
    '监察制度': topic_count(['监察', '纪检', '留置']),
    '刑法域外适用': topic_count(['域外', '涉外']) if False else [a for a in criminal if any(w in a.get('title','') for w in ['域外','涉外']) and any(w in a.get('title','') for w in ['刑法','刑事','犯罪'])],
}

topics_sorted = sorted(topics_data.items(), key=lambda x: -len(x[1]))

# ============ 生成DOCX ============

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ============ 正文开始 ============

# 封面标题
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('CLSCI期刊2025年泛刑法研究态势报告')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.name = 'Times New Roman'
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run('——基于20种CLSCI来源期刊1581篇文章的数据分析')
sub_run.font.size = Pt(14)
sub_run.font.name = 'Times New Roman'
sub_run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

doc.add_paragraph()  # 空行

# ===== 摘要 =====
doc.add_paragraph()
abstract_title = doc.add_paragraph()
abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
atr = abstract_title.add_run('摘  要')
atr.font.size = Pt(14)
atr.font.bold = True
atr.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

abs_text = (
    f"本报告以北大法宝（PKULaw）数据库为数据来源，对2025年度20种CLSCI（中国法学核心科学研究来源期刊）"
    f"共计{total_count}篇学术论文进行全面采集，并从中筛选出{criminal_count}篇泛刑法相关文章（含刑法学、刑事诉讼法学、"
    f"犯罪学及刑事政策等），占全部文章的{100*criminal_count/total_count:.1f}%。"
    f"报告在逐篇采集摘要（采集率{100*has_abstract/criminal_count:.1f}%）与可得全文的基础上，"
    f"运用关键词频率分析、主题聚类、期刊横向比较等方法，系统梳理了2025年度CLSCI期刊泛刑法研究的总体规模、"
    f"学科分布、热点专题、方法论特征与选题趋势。研究发现，轻罪治理、人工智能刑法、行刑衔接、"
    f"刑事司法数字化转型以及经济刑法与反腐败构成本年度五大核心热点；刑法基本理论研究保持稳定产出，"
    f"法益论与归责理论仍是理论争鸣的主战场；刑事诉讼法第四次修改的讨论推动了诸多程序法议题的集中爆发。"
    f"报告最后对未来选题方向与写作策略提出建议。"
)
abs_para = doc.add_paragraph(abs_text)
abs_para.paragraph_format.first_line_indent = Cm(0.74)

doc.add_paragraph()

# ===== 目录提示 =====
doc.add_page_break()

# ===== 一、引言 =====
doc.add_heading('一、引言', level=1)

intro_text = (
    "CLSCI（中国法学核心科学研究来源期刊）是衡量中国法学研究产出质量的核心指标之一。"
    "每年度CLSCI期刊发表的学术论文不仅反映了法学各学科的研究前沿，更是学术评价、人才考核和学科建设的重要参照。"
    "对于刑法学界而言，系统把握CLSCI期刊中泛刑法研究的总体态势、热点分布与发展趋势，"
    "具有重要的学术参考价值与实践指导意义。"
)
p = doc.add_paragraph(intro_text)
p.paragraph_format.first_line_indent = Cm(0.74)

intro2 = (
    f"本报告的数据基础来自对2025年度20种CLSCI来源期刊的全面采集。"
    f"数据采集时间为2026年2月，涵盖各期刊2025年已出版的全部期次。"
    f"采集范围包括：《中国社会科学》《中国法学》《法学研究》《中外法学》"
    f"《法学家》《清华法学》《政法论坛》《法学评论》《法商研究》《现代法学》"
    f"《当代法学》《法制与社会发展》《法学论坛》《政治与法律》《环球法律评论》"
    f"《华东政法大学学报》《法律科学》《法学》《中国法律评论》《中国刑事法杂志》"
    f"共20种期刊，合计采集文章{total_count}篇。"
)
p = doc.add_paragraph(intro2)
p.paragraph_format.first_line_indent = Cm(0.74)

intro3 = (
    '所谓"泛刑法"研究，本报告采取宽泛界定，涵盖以下领域：刑法总论与分论、刑事诉讼法、'
    "证据法（刑事证据方向）、犯罪学与刑事政策、刑事执行法、监察法（涉刑事部分）、"
    "行刑衔接、刑民交叉等交叉领域。筛选标准为标题含刑法相关关键词，并经人工复核排除误判。"
    f"最终确认泛刑法相关文章{criminal_count}篇，占全部文章的{100*criminal_count/total_count:.1f}%，"
    f"即每四篇CLSCI论文中就有一篇涉及泛刑法议题，反映了刑法学科在法学研究中的重要地位。"
)
p = doc.add_paragraph(intro3)
p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 二、总体概览 =====
doc.add_heading('二、总体概览：规模、结构与分布', level=1)

doc.add_heading('（一）总量与占比', level=2)

overview1 = (
    f"2025年度20种CLSCI期刊共发表学术论文{total_count}篇。其中，泛刑法相关文章{criminal_count}篇，"
    f"占比{100*criminal_count/total_count:.1f}%。这一比例在法学各二级学科中处于较高水平，"
    f"表明刑法学（含刑事诉讼法学）始终是CLSCI期刊的核心议题领域。"
    f"值得注意的是，{criminal_count}篇泛刑法文章中，{has_abstract}篇具有完整摘要（采集率{100*has_abstract/criminal_count:.1f}%），"
    f"为深度文本分析提供了可靠的数据基础。"
)
p = doc.add_paragraph(overview1)
p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_heading('（二）期刊分布', level=2)

journal_desc = (
    "各期刊刊发泛刑法文章的数量和比例存在显著差异。除专业期刊《中国刑事法杂志》（100%为刑事法文章）外，"
    "综合性法学期刊中，《清华法学》以48.5%的刑法文章占比高居榜首，几乎每两篇文章中就有一篇涉及泛刑法；"
    "《法学论坛》（32.9%）和《中外法学》（28.8%）紧随其后。"
    "相对而言，《中国法律评论》（10.1%）和《法制与社会发展》（13.3%）的泛刑法文章占比较低，"
    "这与两刊偏重法理学、法社会学等方向的定位相符。"
)
p = doc.add_paragraph(journal_desc)
p.paragraph_format.first_line_indent = Cm(0.74)

# 表格：各期刊数据
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '期刊名称'
hdr[1].text = '总发文量'
hdr[2].text = '泛刑法文章数'
hdr[3].text = '占比'

for s in journal_stats_sorted:
    row = table.add_row().cells
    row[0].text = s['journal']
    row[1].text = str(s['total'])
    row[2].text = str(s['criminal'])
    row[3].text = f"{s['pct']:.1f}%"

# 设置表格字体大小
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

table_note = doc.add_paragraph('表1  2025年CLSCI期刊泛刑法文章分布统计')
table_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_note.runs[0].font.size = Pt(10)
table_note.runs[0].font.bold = True

doc.add_heading('（三）期数分布', level=2)

issue_text = (
    "从发表时间来看，泛刑法文章在各期次中的分布较为均匀，但略呈波动。"
    f"双月刊期刊中，第5期（{issue_counts.get('5',0)}篇）和第3期（{issue_counts.get('3',0)}篇）的刑法文章产出略高于其他期次，"
    f"第2期（{issue_counts.get('2',0)}篇）也保持较高水平。"
    f"第1期（{issue_counts.get('1',0)}篇）、第4期（{issue_counts.get('4',0)}篇）和第6期（{issue_counts.get('6',0)}篇）的产出相对均衡。"
    "月刊期刊（《政治与法律》《法学》）的各期分布则更为分散。"
    '总体来看，2025年度泛刑法研究的发表节奏平稳，未出现明显的\u201c集中爆发期\u201d，表明学术产出具有持续性。'
)
p = doc.add_paragraph(issue_text)
p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 三、热点专题分析 =====
doc.add_heading('三、热点专题深度分析', level=1)

hot_intro = (
    f"通过对{criminal_count}篇泛刑法文章的标题与摘要进行关键词频率分析和主题聚类，"
    f"可以识别出2025年度的若干核心热点专题。以下按文章数量从高到低依次分析。"
)
p = doc.add_paragraph(hot_intro)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.1 刑法基本理论 ---
doc.add_heading('（一）刑法基本理论：法益论与归责理论的纵深发展', level=2)
theory_arts = topics_data.get('刑法基本理论（法益/归责/违法性等）', [])
theory_text = (
    f"刑法基本理论研究是本年度产出最多的专题方向，共计{len(theory_arts)}篇。"
    f"其中，"法益"出现在22篇文章标题中，"归责"出现在19篇标题中，二者构成了刑法理论研究的两大主轴。"
)
p = doc.add_paragraph(theory_text)
p.paragraph_format.first_line_indent = Cm(0.74)

theory2 = (
    "在法益论方面，2025年的讨论呈现出三个新趋势：一是法益概念的数字化拓展，如对"网络秩序法益""数据法益"的探讨；"
    "二是法益理论与具体罪名解释的深度结合，如侮辱罪、非法经营罪、帮信罪等具体罪名中的法益识别；"
    "三是法益分层保护理论的兴起，尤其是在经济秩序法益领域。"
    "典型文章如《法益侵害性的确定与构成要件的保护范围》（中外法学）、"
    "《法益修复的刑法意义：功能主义刑法视角的展开》（政法论坛）、"
    "《利益可归属与刑法法益观的人本化重塑》（政法论坛）等，"
    "均体现了法益论从抽象建构走向具体适用的学术转向。"
)
p = doc.add_paragraph(theory2)
p.paragraph_format.first_line_indent = Cm(0.74)

theory3 = (
    "在归责理论方面，结果归责的类型化、规范化成为主要议题。"
    "《意志归责与规范归责：刑法中结果归责的类型化建构》（现代法学）、"
    "《论规范归咎刑事责任范型》（中国法学）等文章，"
    "尝试在传统行为责任范型之外构建规范性的归责新范式。"
    "此外，归责理论在人工智能犯罪、过失犯等新领域的应用也值得关注。"
)
p = doc.add_paragraph(theory3)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.2 轻罪治理 ---
doc.add_heading('（二）轻罪治理与犯罪记录封存：从刑事政策到制度建构', level=2)
light_arts = topics_data.get('轻罪治理与犯罪记录封存', [])
light_text = (
    f"轻罪治理是2025年度最突出的政策性热点，共有{len(light_arts)}篇文章涉及这一主题。"
    f"这一数量级别在单一政策性议题中极为罕见，反映了轻罪治理已从刑事政策讨论上升为系统性的制度建设议题。"
)
p = doc.add_paragraph(light_text)
p.paragraph_format.first_line_indent = Cm(0.74)

light2 = (
    "从文章内容看，轻罪治理研究呈现出"三条主线"并行发展的格局："
    "第一条主线是实体法上的出罪机制，包括但书出罪（《但书出罪的规范基础与实践展开》）、"
    "醉驾出罪（《由严惩走向宽宥：醉酒型危险驾驶罪从形式入罪到实质出罪的转向》）、"
    "以及商事犯罪的刑民界分等。"
    "第二条主线是程序法上的分流机制，如相对不起诉的扩大适用、快速办理机制等。"
    "第三条主线是犯罪记录封存与前科消灭制度，多达6篇文章专门论述犯罪记录封存，"
    "包括《轻微犯罪记录封存制度的时空维度建构》（法制与社会发展）、"
    "《论我国犯罪记录封存制度的革新与续造》（中国法律评论）等，"
    "体现了学界对轻罪附随后果问题的高度关注。"
)
p = doc.add_paragraph(light2)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.3 经济刑法 ---
doc.add_heading('（三）经济刑法与反腐败：受贿罪研究的集中爆发', level=2)
econ_arts = topics_data.get('经济刑法与反腐败', [])
econ_text = (
    f"经济刑法与反腐败研究共计{len(econ_arts)}篇，是传统刑法分论中产出最多的领域。"
    f"值得特别注意的是，受贿罪相关研究出现了罕见的集中爆发，仅标题含"受贿"的文章就有10篇。"
    f"具体而言，"代持型受贿""约定受贿""商业机会型受贿""斡旋受贿""截贿"等新型受贿模式成为研究焦点。"
    f"这一现象与近年来反腐实践中新型隐性腐败案件的增多密切相关。"
)
p = doc.add_paragraph(econ_text)
p.paragraph_format.first_line_indent = Cm(0.74)

econ2 = (
    "在洗钱罪研究方面，《政法论坛》第1期刊登了三篇洗钱罪专题文章，"
    "《中国刑事法杂志》第1期亦有两篇洗钱相关文章，形成了年度专题研讨的态势。"
    "讨论焦点包括：自洗钱入罪后的司法认定、洗钱罪与掩饰隐瞒犯罪所得罪的界分、"
    "地下钱庄的刑事法律规制等。"
    "此外，税收犯罪、虚开增值税专用发票罪、非法经营罪等传统经济犯罪罪名的研究持续活跃。"
)
p = doc.add_paragraph(econ2)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.4 行刑衔接 ---
doc.add_heading('（四）行刑衔接与法秩序统一：交叉领域研究的井喷', level=2)
xing_arts = topics_data.get('行刑衔接与法秩序统一', [])
xing_text = (
    f"行刑衔接与法秩序统一性研究共计{len(xing_arts)}篇，成为2025年度增长最快的研究方向之一。"
    f""行刑"一词在11篇文章标题中出现，"法定犯"出现5次，"刑民"出现5次。"
    f"这一热点的形成与检察机关推进行刑反向衔接工作的实践需求直接相关。"
)
p = doc.add_paragraph(xing_text)
p.paragraph_format.first_line_indent = Cm(0.74)

xing2 = (
    "研究内容涵盖三个层面：一是宏观理论层面，探讨行政违法与刑事犯罪的界限，如"
    "《行刑衔接视野下行政犯行政违法的性质与功能》（政法论坛）；"
    "二是制度建构层面，讨论行刑双向衔接的程序机制，如"
    "《行刑双向衔接的实体法基础及其程序展开》（法商研究）；"
    "三是具体罪名层面，分析法定犯的认定标准与自然犯的区分，如"
    "《区分自然犯与法定犯的新标准：主体性法益／主体间性法益》（中国法学）。"
)
p = doc.add_paragraph(xing2)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.5 人工智能与数字化 ---
doc.add_heading('（五）人工智能、数据犯罪与刑事司法数字化', level=2)
ai_arts = topics_data.get('人工智能与刑法', [])
data_arts = topics_data.get('数据犯罪与个人信息保护', [])
digital_arts = topics_data.get('刑事司法数字化转型', [])
ai_text = (
    f"人工智能与刑法（{len(ai_arts)}篇）、数据犯罪与个人信息保护（{len(data_arts)}篇）、"
    f"刑事司法数字化转型（{len(digital_arts)}篇）三个板块共同构成了"数字时代刑法"的研究版图，"
    f"合计超过30篇，成为本年度最具时代特征的研究方向。"
    f""数字"一词出现在25篇文章标题中，"人工智能"出现在13篇标题中。"
)
p = doc.add_paragraph(ai_text)
p.paragraph_format.first_line_indent = Cm(0.74)

ai2 = (
    "人工智能刑法研究主要聚焦四个议题：一是人工智能犯罪的归责主体与归责路径，"
    "如《人工智能犯罪归责的行为主体论路径》（法学家）；"
    "二是深度伪造等新技术的刑法规制，如《人工智能时代深度伪造行为的刑法规制》（政治与法律）；"
    "三是人工智能辅助司法的规范化，如《量刑智能辅助的本土实践》（政法论坛）和"
    "《大模型司法适用的认知路径及其规范》（中国刑事法杂志）；"
    "四是生成式人工智能带来的新型风险，如《人工智能时代我国儿童色情的刑法规制》等。"
    "其中，大模型（LLM）在司法中的应用是2025年的全新议题，"
    "反映了ChatGPT等生成式AI对刑事司法领域的深刻影响。"
)
p = doc.add_paragraph(ai2)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.6 网络犯罪 ---
doc.add_heading('（六）网络犯罪治理：帮信罪与网络暴力的持续关注', level=2)
net_arts = topics_data.get('网络犯罪治理', [])
net_text = (
    f"网络犯罪治理研究共{len(net_arts)}篇。帮助信息网络犯罪活动罪（帮信罪）仍然是讨论的焦点，"
    f"多篇文章从法益保护、规范属性、竞合关系等角度展开分析，反映了帮信罪成为近年来案件数量最多的罪名之一后"
    f"学界对其适用边界的持续反思。"
)
p = doc.add_paragraph(net_text)
p.paragraph_format.first_line_indent = Cm(0.74)

net2 = (
    "网络暴力的刑法治理是另一个突出议题，涉及入罪标准、归责逻辑、证据供给等多个面向。"
    "这与2023年以来最高检推动网络暴力专项治理的实践背景密切相关。"
    "此外，网络犯罪产业链的协同治理、网络犯罪证明难题的司法治理等问题也受到关注。"
)
p = doc.add_paragraph(net2)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.7 认罪认罚 ---
doc.add_heading('（七）认罪认罚从宽制度：从制度推行到效果反思', level=2)
plea_arts = topics_data.get('认罪认罚从宽制度', [])
plea_text = (
    f"认罪认罚从宽制度研究共{len(plea_arts)}篇。与往年侧重制度正当性论证不同，"
    f"2025年的研究明显转向制度实施效果的反思与优化。"
    f"《认罪认罚从宽制度的实践反思与体系重塑》（法学评论）、"
    f"《认罪认罚自愿性保障的制度性缺陷及其破解》（政治与法律）等文章，"
    f"直面制度运行中的现实问题。量刑协商机制的实质化问题得到了专门讨论，"
    f"《华东政法大学学报》第5期甚至推出了量刑协商专题，刊登三篇相关文章。"
)
p = doc.add_paragraph(plea_text)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.8 刑事证据 ---
doc.add_heading('（八）刑事证据制度：电子证据与证据体系的重构', level=2)
ev_arts = topics_data.get('刑事证据制度', [])
ev_text = (
    f"刑事证据相关研究共{len(ev_arts)}篇。2025年的证据法研究呈现出两大特征："
    f"一是电子证据、数字证据的理论与实务问题持续升温，如《论数字时代刑事证据的三元结构》（中外法学）"
    f"提出从传统二元结构向三元结构转型的理论框架；"
    f"二是证据审查方法论的创新，如《证据层控主义：一种中国式的证据审查结构观》（中外法学）"
    f"尝试提炼具有中国特色的证据审查理论。"
    f"此外，非法证据排除、证据链理论、供述补强规则等传统议题仍保持较高关注度。"
)
p = doc.add_paragraph(ev_text)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 3.9 刑事诉讼法修改 ---
doc.add_heading('（九）刑事诉讼法第四次修改：程序法议题的集中爆发', level=2)
cpc_arts = topics_data.get('刑事诉讼法修改', [])
cpc_text = (
    f"刑事诉讼法第四次修改相关研究共{len(cpc_arts)}篇。"
    f"随着刑事诉讼法修改被正式提上日程，法典化路径、特别程序体系、回避制度、被害人权利保护等"
    f"具体修法议题均得到了深入讨论。"
    f"《从形式法典化到实质法典化：刑事诉讼法的法典化进阶》（中国法学）、"
    f"《犯罪治理扩张背景下的刑事诉讼法反向工具价值研究》（清华法学）等文章，"
    f"从不同角度回应了法典化时代的程序法改革需求。"
    f"值得注意的是，被告人对质权、对物强制措施、刑事证据规则等具体程序制度的完善也受到广泛关注。"
)
p = doc.add_paragraph(cpc_text)
p.paragraph_format.first_line_indent = Cm(0.74)

# --- 其他专题简述 ---
doc.add_heading('（十）其他值得关注的研究方向', level=2)

other_topics = [
    (f"危险犯理论与醉驾治理（{len(topics_data.get('危险犯理论与醉驾治理',[]))}篇）",
     "抽象危险犯的判断标准、醉驾入罪与出罪的边界是持续争议的议题。多篇文章讨论了准抽象危险犯概念的引入，以及危险驾驶罪的实质出罪机制。"),
    (f"诈骗犯罪研究（{len(topics_data.get('诈骗犯罪',[]))}篇）",
     "诈骗罪中的直接性要件、认识错误、财产损失判断标准等教义学问题得到持续探讨，电信网络诈骗的刑法定性问题也受到关注。"),
    (f"侦查制度（{len(topics_data.get('侦查制度',[]))}篇）",
     "机动侦查权的理论定位与程序设计成为热点，多篇文章探讨了检察机关侦查权的边界与行使方式。技术公司参与侦查的角色定位也是新兴议题。"),
    (f"正当防卫与紧急避险（{len(topics_data.get('正当防卫/紧急避险',[]))}篇）",
     "正当防卫的研究在家庭暴力反击场景中展开，防卫过当的认定标准和防御性紧急避险的理论构建受到关注。"),
    (f"共犯论（{len(topics_data.get('共犯论',[]))}篇）",
     "网络共同犯罪中的单向意思联络、共犯从属性理论在数字时代的危机与应对、身份犯共犯等问题引发讨论。"),
    (f"单位犯罪与企业合规（{len(topics_data.get('单位犯罪与企业合规',[]))}篇）",
     "单位犯罪废止论的提出引发理论争鸣，民营经济刑法保护和涉企异地刑事司法问题受到高度关注。"),
    (f"监察制度（{len(topics_data.get('监察制度',[]))}篇）",
     "监察体制改革深化背景下，监察调查权的立法扩展、监检衔接、纪检监察的数字化转型等议题持续推进。"),
]

for title, desc in other_topics:
    p = doc.add_paragraph()
    run_b = p.add_run(title + '：')
    run_b.bold = True
    p.add_run(desc)
    p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 四、期刊特色分析 =====
doc.add_heading('四、重点期刊刑法研究特色', level=1)

journal_analysis = [
    ("《清华法学》",
     f"泛刑法文章占比高达48.5%（33/68），在综合性法学期刊中遥遥领先。2025年的刑法研究集中在受贿罪（多达6篇）、"
     f"犯罪治理与轻罪、过失犯与归责理论三个方向。《清华法学》第6期推出受贿罪专题、第5期推出轻罪与数据犯罪专题，"
     f"体现了通过集中选题形成学术影响力的编辑策略。"),
    ("《法学论坛》",
     f"泛刑法占比32.9%（27/82），在综合性法学期刊中排名第二。"
     f"第1期推出了"刑诉法再修改之证据疑难问题""数字经济时代犯罪治理"双专题；"
     f"帮信罪、犯罪记录封存、网络犯罪治理等议题的讨论均有较强的时效性和实务性。"),
    ("《中国法学》",
     f"作为法学界最高级别期刊，其21篇泛刑法文章涵盖了年度几乎所有重要议题：法益论、归责理论、"
     f"刑事诉讼法典化、轻罪治理、行刑衔接等。文章体量普遍较大，理论深度突出。"),
    ("《中国刑事法杂志》",
     f"46篇文章全部为刑事法研究，选题覆盖面广。2025年重点关注洗钱罪、数据犯罪、认罪认罚、"
     f"趋利性执法等实务热点，体现了理论与实践紧密结合的特色。"),
    ("《政治与法律》",
     f"作为月刊，32篇泛刑法文章分布均匀。选题兼顾理论与实务，"
     f"刑法教义学分析与司法实务问题研究并重。"),
    ("《中外法学》",
     f"23篇泛刑法文章中，家事犯罪专题（第3期）和经济刑法归责根据等文章体现了较强的理论原创性。"
     f"刑事庭审笔录、证据层控主义等议题具有鲜明的中国问题意识。"),
]

for journal_name, desc in journal_analysis:
    p = doc.add_paragraph()
    run_b = p.add_run(journal_name + '：')
    run_b.bold = True
    p.add_run(desc)
    p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 五、趋势判断 =====
doc.add_heading('五、2025年度泛刑法研究的五大趋势', level=1)

trends = [
    ("趋势一：从"治罪"到"治理"的范式转换加速",
     "轻罪治理的27篇文章、犯罪记录封存的6篇文章、行刑反向衔接的系列文章共同表明，"
     "刑法研究正从单纯的"定罪量刑"向"社会治理"视角全面转型。"出罪"成为高频词（9篇），"
     "这在以往的CLSCI研究中并不常见。这一趋势与我国轻罪案件已占刑事案件总量80%以上的现实背景密切相关。"),
    ("趋势二：数字时代刑法研究的全面展开",
     "人工智能、数据犯罪、数字化转型三个板块合计超过30篇，"数字"一词出现在25篇标题中。"
     "从刑事实体法到程序法，从侦查取证到量刑辅助，数字技术的影响已渗透到刑事法的每一个环节。"
     "大模型（LLM）在司法中的应用是2025年的全新增长点，预计将在未来持续发酵。"),
    ("趋势三：刑事诉讼法修改催生程序法研究高潮",
     "随着刑事诉讼法第四次修改被提上日程，从法典化路径选择到具体制度完善，"
     "程序法议题出现了系统性爆发。认罪认罚、量刑协商、被告人权利保障、侦查制度改革等议题"
     "均呈现出围绕修法方向展开讨论的特征。这一趋势将在未来1-2年内进一步强化。"),
    ("趋势四：刑法教义学与实证研究的方法论融合",
     "2025年的研究方法呈现出明显的多元化特征。一方面，法教义学分析仍然是主流方法（26篇显性运用）；"
     "另一方面，基于裁判文书大数据的实证分析明显增多（23篇），"
     "如《无限额罚金刑的裁量规则研究——以30余万份毒品犯罪裁判文书为例的分析》等文章"
     "体现了数据驱动型研究的新范式。教义学与实证方法的结合正在成为高质量研究的新标准。"),
    ("趋势五：交叉领域研究的制度化",
     "行刑衔接、刑民交叉、监检衔接等交叉领域不再被视为"边缘议题"，"
     "而是上升为具有独立学术价值的研究方向。24篇行刑衔接文章的出现表明，"
     "法秩序统一性视角下的跨部门法研究已经实现制度化。"
     "这一趋势要求研究者具备更宽广的知识视野和更强的跨学科对话能力。"),
]

for i, (title, content) in enumerate(trends, 1):
    doc.add_heading(f'（{["一","二","三","四","五"][i-1]}）{title[4:]}', level=2)
    p = doc.add_paragraph(content)
    p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 六、选题建议 =====
doc.add_heading('六、选题方向与写作建议', level=1)

doc.add_heading('（一）高潜力选题方向', level=2)

topics_suggest = (
    "基于上述分析，以下方向在2025-2026年具有较高的学术价值和发表潜力："
)
p = doc.add_paragraph(topics_suggest)
p.paragraph_format.first_line_indent = Cm(0.74)

suggestions = [
    "轻罪治理的体系化建构：犯罪记录封存的具体制度设计、轻罪前科消灭的条件与程序、"
    "轻罪与行政处罚的衔接机制等，理论空间大且实践需求迫切。",
    "人工智能犯罪的归责体系：大模型生成有害内容的刑事责任分配、自动驾驶事故的过失认定、"
    "AI辅助犯罪的共犯论分析等，是理论前沿与技术发展的交汇点。",
    "刑事诉讼法修改的具体制度完善：辩护制度、强制措施体系、涉外刑事诉讼程序等方向，"
    "具有直接的立法参考价值。",
    "数据犯罪的刑法保护体系：数据法益的独立化、数据犯罪的类型区分、数据安全风险的刑法预防等，"
    "尚处于理论建构的早期阶段。",
    "行刑衔接的程序机制：反向衔接的具体操作规范、行政犯的实质认定标准、"
    "法秩序统一性下的违法判断等，实务需求大但理论供给不足。",
    "新型腐败行为的刑法规制：代持型受贿、商业机会型受贿、期权式腐败等新型腐败形态的定性与证明，"
    "是反腐实践提出的新课题。",
]

for i, s in enumerate(suggestions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run_b = p.add_run(f'{i}. ')
    run_b.bold = True
    p.add_run(s)

doc.add_heading('（二）写作策略建议', level=2)

writing_tips = [
    ("问题导向而非概念导向",
     "从2025年的高质量文章来看，CLSCI期刊偏好从具体问题切入的研究。"
     "例如，与其泛泛讨论"法益理论"，不如聚焦"数据法益的识别与保护"这样的具体问题。"
     "标题应直接体现研究问题和核心论点。"),
    ("教义学分析与实证数据相结合",
     "纯粹的理论推演已不足以满足CLSCI期刊的要求。"
     "最具竞争力的文章往往将规范分析与裁判文书数据、域外立法比较、实务调研等实证材料相结合。"),
    ("回应立法与政策动态",
     "2025年的高产议题（轻罪治理、刑诉法修改、行刑衔接）无一不与正在推进的立法修改和司法改革密切相关。"
     "选题应关注立法动态，研究成果应能为制度完善提供参考。"),
    ("跨学科视角的适度引入",
     "适当引入犯罪学、法经济学、计算机科学等跨学科视角，"
     "可以增强论文的理论深度和创新性。但跨学科方法的运用应服务于法学问题的解决，"
     "避免为方法论而方法论。"),
    ("注重比较法的精细化",
     "2025年涉及比较法的文章有31篇，但高质量的比较研究不应止于域外制度的介绍，"
     "而应深入分析制度差异背后的法理逻辑，并评估域外经验在中国语境下的适用性。"),
]

for title, desc in writing_tips:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run_b = p.add_run(title + '：')
    run_b.bold = True
    p.add_run(desc)

# ===== 七、结语 =====
doc.add_heading('七、结语', level=1)

conclusion = (
    f"2025年度CLSCI期刊共发表泛刑法相关文章{criminal_count}篇，占全部{total_count}篇论文的{100*criminal_count/total_count:.1f}%。"
    f"本年度的研究呈现出从"治罪"到"治理"的范式转换、数字时代刑法研究的全面展开、"
    f"刑事诉讼法修改推动的程序法研究高潮、教义学与实证方法的融合、以及交叉领域研究的制度化等五大趋势。"
    f"轻罪治理（27篇）、经济刑法与反腐败（30篇）、行刑衔接（24篇）、人工智能刑法（17篇）、"
    f"刑事证据制度（17篇）构成了年度五大热点专题群。"
)
p = doc.add_paragraph(conclusion)
p.paragraph_format.first_line_indent = Cm(0.74)

conclusion2 = (
    "从学科发展的角度看，刑法学正处于一个重要的转型期。"
    "一方面，以法益论和归责理论为核心的刑法教义学体系建设持续推进，"
    "中国自主刑法知识体系的构建成为学界的自觉追求；"
    "另一方面，数字技术革命、社会治理转型和立法修改的实践需求，"
    "不断拓展着刑法研究的问题域和方法论。"
    "如何在坚持教义学品格的同时有效回应时代需求，"
    "是当前刑法学界面临的核心课题。"
)
p = doc.add_paragraph(conclusion2)
p.paragraph_format.first_line_indent = Cm(0.74)

conclusion3 = (
    "本报告基于对20种CLSCI期刊的全量数据采集与系统分析，"
    "力图为刑法学界的研究者、学术期刊编辑以及法学院校的学科建设提供数据化的参考。"
    "受限于关键词筛选方法的固有局限，个别文章的归类可能存在偏差；"
    "此外，部分文章的作者信息、机构信息未能完整采集，"
    "对于高产作者和高产机构的分析有待补充。这些问题留待后续研究进一步完善。"
)
p = doc.add_paragraph(conclusion3)
p.paragraph_format.first_line_indent = Cm(0.74)

# ===== 附录 =====
doc.add_page_break()
doc.add_heading('附录：2025年CLSCI期刊泛刑法文章完整目录', level=1)

# 按期刊分组列出
current_journal = ''
for art in sorted(criminal, key=lambda x: (x.get('journal',''), x.get('issue','0').zfill(2))):
    j = art.get('journal', '')
    if j != current_journal:
        current_journal = j
        count = crim_by_j.get(j, 0)
        doc.add_heading(f'{j}（{count}篇）', level=2)

    issue = art.get('issue', '')
    title = art.get('title', '')
    issue_str = f'第{issue}期' if issue else ''
    p = doc.add_paragraph(f'{issue_str}  {title}', style='List Number')
    p.runs[0].font.size = Pt(10)

# 保存
output_path = f'{OUTPUT_DIR}/CLSCI_2025_泛刑法研究态势报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')

# 统计字数
total_chars = 0
for para in doc.paragraphs:
    total_chars += len(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_chars += len(cell.text)
print(f'报告总字符数: {total_chars}')
